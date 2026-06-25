import re
from io import BytesIO
from datetime import datetime
from typing import Union
import docx

def calculate_age(dob_str: str) -> str:
    """
    Calculates age based on Date of Birth string.
    Supports formats: DD/MM/YYYY, YYYY-MM-DD, DD-MM-YYYY, YYYY/MM/DD.
    Falls back to simple birth-year subtraction if format is unparseable.
    """
    if not dob_str:
        return ""
    
    # Try common datetime formats
    for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%Y/%m/%d"):
        try:
            dob_dt = datetime.strptime(dob_str.strip(), fmt)
            today = datetime.today()
            age = today.year - dob_dt.year - ((today.month, today.day) < (dob_dt.month, dob_dt.day))
            return str(age)
        except ValueError:
            continue
            
    # Regex fallback to find a 4-digit year (e.g. 1990)
    years = re.findall(r"\b(19\d\d|20\d\d)\b", dob_str)
    if years:
        dob_year = int(years[0])
        return str(datetime.today().year - dob_year)
        
    return ""

def apply_acceptance_rules(data: dict) -> dict:
    """
    Automatically populates the 'Acceptance Conditions' for each member block
    based on the presence of exclusions.
    - Clean case: 'Accepted'
    - Case with exclusions: 'Accepted with exclusion'
    """
    updated = data.copy()
    blocks = [
        ("name", "exclusions", "acceptance_conditions"),
        ("sp_name", "sp_exclusions", "sp_acceptance_conditions"),
        ("c1_name", "c1_exclusions", "c1_acceptance_conditions"),
        ("c2_name", "c2_exclusions", "c2_acceptance_conditions"),
        ("c3_name", "c3_exclusions", "c3_acceptance_conditions"),
    ]
    
    # Values of exclusions that indicate a "clean" case (no exclusions)
    clean_indicators = {
        "", "none", "none.", "n/a", "na", "no", "no exclusion", 
        "no exclusions", "-", "clean", "nil", "no pre-existing conditions"
    }
    
    for name_key, excl_key, status_key in blocks:
        name_val = str(updated.get(name_key) or "").strip()
        if name_val:
            excl_val = str(updated.get(excl_key) or "").strip()
            # Remove trailing/leading spaces and lowercase for check
            check_val = excl_val.lower().rstrip(".")
            if not check_val or check_val in clean_indicators:
                updated[status_key] = "Accepted"
            else:
                updated[status_key] = "Accepted with exclusion"
        else:
            # If the person doesn't exist, clear acceptance conditions
            updated[status_key] = ""
            
    return updated

def resolve_plan_combination(data: dict, config: dict = None) -> dict:
    """
    Resolves plan codes and formats the plan and deductible fields.
    """
    import re
    if config is None:
        from src.pdf_processor.inverter import load_config_by_pdf_id
        pdf_id = data.get("pdf_id")
        config = load_config_by_pdf_id(pdf_id)

    products = config.get("product_options", {}).get("products", {})
    product_keys = list(products.keys())
    
    prefix_a = "ESSENTIAL"
    prefix_b = "VISA"
    
    if len(product_keys) >= 2:
        product_name_a = product_keys[0]
        choices_a = products[product_name_a].get("plan_tier", {}).get("choices", [])
        if choices_a:
            prefix_a = choices_a[0].rstrip("0123456789")
            
        product_name_b = product_keys[1]
        choices_b = products[product_name_b].get("plan_tier", {}).get("choices", [])
        if choices_b:
            prefix_b = choices_b[0].rstrip("0123456789")

    updated = data.copy()
    plan_val = str(updated.get("plan") or "").strip()
    ded_val = str(updated.get("deductible") or "").strip()
    
    if not updated.get("product_name") and plan_val:
        if prefix_a.upper() in plan_val.upper() or "SMARTCARE" in plan_val.upper():
            choices = config.get("product_options", {}).get("product_name", {}).get("choices", ["ESSENTIAL"])
            updated["product_name"] = choices[0]
        elif prefix_b.upper() in plan_val.upper() or "EASYCARE" in plan_val.upper():
            choices = config.get("product_options", {}).get("product_name", {}).get("choices", ["ESSENTIAL", "EASYCARE"])
            updated["product_name"] = choices[1] if len(choices) > 1 else choices[0]

    if not plan_val:
        return updated
        
    plan_tier = ""
    optional_benefit = ""
    opd_choice = ""
    is_visa = False
    
    # Check if VISA or prefix_b
    if prefix_b.upper() in plan_val.upper() or "VISA" in plan_val.upper():
        is_visa = True
        match = re.search(rf"(?:{prefix_b}|VISA)\s*(?:Plan)?\s*(\d+)", plan_val, re.IGNORECASE)
        if match:
            plan_tier = match.group(1)
        else:
            match = re.search(r"Plan\s*(\d+)", plan_val, re.IGNORECASE)
            if match:
                plan_tier = match.group(1)
    else:
        # prefix_a or ESSENTIAL
        match_ess = re.search(rf"(?:{prefix_a}|ESSENTIAL)\s*(\d+)", plan_val, re.IGNORECASE)
        if match_ess:
            plan_tier = match_ess.group(1)
        else:
            match_plan = re.search(r"Plan\s*(\d+)", plan_val, re.IGNORECASE)
            if match_plan:
                plan_tier = match_plan.group(1)

        # Extract benefit: IPD+OPD+WELLNESS or IPD+OPD or IPD
        if "IPD+OPD+WELLNESS" in plan_val:
            optional_benefit = "IPD+OPD+WELLNESS"
        elif "IPD+OPD" in plan_val:
            optional_benefit = "IPD+OPD"
        elif "IPD" in plan_val:
            optional_benefit = "IPD"

        # Extract OPD choice (inside parentheses or after dash/space)
        match_opd = re.search(r"\(([^)]+)\)", plan_val)
        if match_opd:
            opd_choice = match_opd.group(1).strip()
        else:
            parts = [p.strip() for p in plan_val.split("-")]
            for p in parts:
                if p in ("3k * 30 times / year", "50k per year"):
                    opd_choice = p
                    
    deductible_amount = ded_val.replace("k", ",000").replace("K", ",000")
    if deductible_amount == "0,000":
        deductible_amount = "0"
         
    if plan_tier:
        if is_visa:
            combo_key = f"{prefix_b}{plan_tier} DD {deductible_amount}"
        elif optional_benefit:
            if optional_benefit == "IPD":
                combo_key = f"{prefix_a}{plan_tier}-IPD DD {deductible_amount}"
            else:
                combo_key = f"{prefix_a}{plan_tier}-{optional_benefit}({opd_choice}) DD {deductible_amount}"
        else:
            return updated
            
        try:
            combo_map = config.get("combinations_map", {})
            if combo_key not in combo_map:
                if prefix_a in combo_key:
                    alt = combo_key.replace(prefix_a, "MOCKA" if prefix_a == "ESSENTIAL" else "ESSENTIAL")
                    if alt in combo_map:
                        combo_key = alt
                elif prefix_b in combo_key:
                    alt = combo_key.replace(prefix_b, "MOCKB" if prefix_b == "VISA" else "VISA")
                    if alt in combo_map:
                        combo_key = alt
            
            plan_code = combo_map.get(combo_key)
            if plan_code:
                updated["plan"] = f"{combo_key} ({plan_code})"
                updated["deductible"] = deductible_amount
        except Exception:
            pass
            
    return updated



def fill_blue_table_docx(template: Union[str, BytesIO], data: dict, config: dict = None) -> BytesIO:
    """
    Fills the BlueTable.docx template tables with the provided data dict.
    Returns the filled file as a BytesIO stream.
    """
    data = resolve_plan_combination(data, config=config)
    data = apply_acceptance_rules(data)
    doc = docx.Document(template)
    
    # Table 0: Main Insured
    age = data.get("age", "")
    if not age and data.get("dob"):
        age = calculate_age(data["dob"])
        
    policy_ver = data.get("policy_version", "")
    policy_ver_val = ""
    if policy_ver:
        if str(policy_ver).lower() in ("thai", "th"):
            policy_ver_val = "TH"
        elif str(policy_ver).lower() in ("english", "en"):
            policy_ver_val = "EN"
        else:
            policy_ver_val = str(policy_ver)

    t0_mapping = {
        "Main Insured": data.get("name", ""),
        "Date of Birth": data.get("dob", ""),
        "Age": age,
        "ID No./Passport No.": data.get("id_card_no", ""),
        "Nationality": data.get("nationality", ""),
        "Beneficiary name": data.get("beneficiary", ""),
        "Relation": data.get("bene_relation", ""),
        "Occupation": data.get("occupation", ""),
        "Agent CODE/Name": data.get("agent", ""),
        "Product Name": data.get("product_name", ""),
        "Policy Version": policy_ver_val,
        "Plan": data.get("plan", ""),
        "Deductible": data.get("deductible", ""),
        "Premium": data.get("premium", ""),
        "Effective date": data.get("effective_date", ""),
        "Personal Address": data.get("present_address", ""),
        "Tel": data.get("tel", ""),
        "Email": data.get("email", ""),
        "Payor Name": data.get("payor_name", ""),
        "Payor Address": data.get("payor_address", ""),
        "TAX ID": data.get("tax_id", ""),
        "Acceptance Conditions": data.get("acceptance_conditions", ""),
        "Exclusions": data.get("exclusions", ""),
    }

    # Table 1: Spouse
    sp_age = data.get("sp_age", "")
    if not sp_age and data.get("sp_dob"):
        sp_age = calculate_age(data["sp_dob"])
        
    t1_mapping = {
        "Spouse": data.get("sp_name", ""),
        "Date of Birth": data.get("sp_dob", ""),
        "Age": sp_age,
        "ID No./Passport No.": data.get("sp_id_card_no", ""),
        "Nationality": data.get("sp_nationality", ""),
        "Beneficiary name": data.get("sp_beneficiary", ""),
        "Relation": data.get("sp_bene_relation", ""),
        "Occupation": data.get("sp_occupation", ""),
        "Acceptance Conditions": data.get("sp_acceptance_conditions", ""),
        "Exclusions": data.get("sp_exclusions", ""),
    }

    # Table 2: Child 1
    c1_age = data.get("c1_age", "")
    if not c1_age and data.get("c1_dob"):
        c1_age = calculate_age(data["c1_dob"])
        
    t2_mapping = {
        "Child # 1": data.get("c1_name", ""),
        "Date of Birth": data.get("c1_dob", ""),
        "Age": c1_age,
        "ID No./Passport No.": data.get("c1_id_card_no", ""),
        "Nationality": data.get("c1_nationality", ""),
        "Beneficiary name": data.get("c1_beneficiary", ""),
        "Relation": data.get("c1_bene_relation", ""),
        "Occupation": data.get("c1_occupation", ""),
        "Acceptance Conditions": data.get("c1_acceptance_conditions", ""),
        "Exclusions": data.get("c1_exclusions", ""),
    }

    # Table 3: Child 2
    c2_age = data.get("c2_age", "")
    if not c2_age and data.get("c2_dob"):
        c2_age = calculate_age(data["c2_dob"])
        
    t3_mapping = {
        "Child # 2": data.get("c2_name", ""),
        "Date of Birth": data.get("c2_dob", ""),
        "Age": c2_age,
        "ID No./Passport No.": data.get("c2_id_card_no", ""),
        "Nationality": data.get("c2_nationality", ""),
        "Beneficiary name": data.get("c2_beneficiary", ""),
        "Relation": data.get("c2_bene_relation", ""),
        "Occupation": data.get("c2_occupation", ""),
        "Acceptance Conditions": data.get("c2_acceptance_conditions", ""),
        "Exclusions": data.get("c2_exclusions", ""),
    }

    # Table 4: Child 3
    c3_age = data.get("c3_age", "")
    if not c3_age and data.get("c3_dob"):
        c3_age = calculate_age(data["c3_dob"])
        
    t4_mapping = {
        "Child # 3": data.get("c3_name", ""),
        "Date of Birth": data.get("c3_dob", ""),
        "Age": c3_age,
        "ID No./Passport No.": data.get("c3_id_card_no", ""),
        "Nationality": data.get("c3_nationality", ""),
        "Relation": data.get("c3_bene_relation", ""),
        "Occupation": data.get("c3_occupation", ""),
        "Acceptance Conditions": data.get("c3_acceptance_conditions", ""),
        "Exclusions": data.get("c3_exclusions", ""),
    }

    mappings = [t0_mapping, t1_mapping, t2_mapping, t3_mapping, t4_mapping]

    def set_cell_text(cell, text):
        p = cell.paragraphs[0]
        p.text = str(text) if text is not None else ""
        # Remove extra paragraphs in the cell to avoid adding unintended newlines
        for extra_p in cell.paragraphs[1:]:
            p_element = extra_p._p
            p_element.getparent().remove(p_element)

    # Populate matching row keys
    for i, table in enumerate(doc.tables):
        if i >= len(mappings):
            break
        mapping = mappings[i]
        
        for row in table.rows:
            col0_text = row.cells[0].text.strip()
            if "General Conditions" in col0_text:
                prod_name = str(data.get("product_name") or "").upper()
                
                # Look up general conditions from config
                try:
                    if config is None:
                        from src.pdf_processor.inverter import load_config_by_pdf_id
                        pdf_id = data.get("pdf_id")
                        config = load_config_by_pdf_id(pdf_id)
                        
                    products = config.get("product_options", {}).get("products", {})
                    product_keys = list(products.keys())
                    product_name_a = product_keys[0] if len(product_keys) >= 1 else "SmartCare Essential"
                    product_name_b = product_keys[1] if len(product_keys) >= 2 else "EasyCare Visa"
                    
                    choices_b = products[product_name_b].get("plan_tier", {}).get("choices", []) if len(product_keys) >= 2 else []
                    prefix_b = choices_b[0].rstrip("0123456789") if choices_b else "VISA"
                    
                    resolved_prod_name = product_name_a
                    if prefix_b.upper() in prod_name or "EASYCARE" in prod_name or "VISA" in prod_name or "MOCKB" in prod_name:
                        row.cells[0].text = f"General Conditions:\n{product_name_b}\n"
                        resolved_prod_name = product_name_b
                    else:
                        row.cells[0].text = f"General Conditions:\n{product_name_a}\n"
                        resolved_prod_name = product_name_a
                        
                    gen_conds = config.get("general_conditions", {})
                    prod_rules = gen_conds.get(resolved_prod_name, {})
                    
                    policy_ver = str(data.get("policy_version") or "").lower()
                    if policy_ver in ("thai", "th"):
                        target_lang = "Thai"
                    elif policy_ver in ("english", "en"):
                        target_lang = "English"
                    else:
                        target_lang = "Both"
                        
                    resolved_msg = prod_rules.get(target_lang)
                    if not resolved_msg:
                        for fallback_lang in (target_lang, "Both", "Thai", "English"):
                            for k, v in prod_rules.items():
                                if k.lower() == fallback_lang.lower():
                                    resolved_msg = v
                                    break
                            if resolved_msg:
                                break
                    if not resolved_msg and prod_rules:
                        resolved_msg = list(prod_rules.values())[0]
                        
                    if resolved_msg:
                        set_cell_text(row.cells[1], resolved_msg)
                except Exception:
                    set_cell_text(row.cells[1], "[General Conditions Placeholder]")

            if col0_text in mapping:
                set_cell_text(row.cells[1], mapping[col0_text])

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
