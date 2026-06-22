import sys
from pathlib import Path
import pytest
import grpc
import resource
from unittest.mock import patch
from concurrent import futures
import time

# Setup paths
worker_dir = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(worker_dir))
sys.path.insert(0, str(worker_dir / "src"))

from proto import document_pb2
from proto import document_pb2_grpc
from src.server import DocumentServiceServicer, configure_resource_limits

@pytest.fixture(scope="module")
def grpc_server():
    # Set max_send_message_length and max_receive_message_length to 10MB to allow testing 5MB limit
    options = [
        ('grpc.max_send_message_length', 10 * 1024 * 1024),
        ('grpc.max_receive_message_length', 10 * 1024 * 1024)
    ]
    server = grpc.server(futures.ThreadPoolExecutor(max_workers=1), options=options)
    document_pb2_grpc.add_DocumentServiceServicer_to_server(
        DocumentServiceServicer(), server
    )
    port = server.add_insecure_port('[::]:0') # Use any available port
    server.start()
    yield f'localhost:{port}'
    server.stop(0)

@pytest.fixture(scope="module")
def grpc_stub(grpc_server):
    options = [
        ('grpc.max_send_message_length', 10 * 1024 * 1024),
        ('grpc.max_receive_message_length', 10 * 1024 * 1024)
    ]
    with grpc.insecure_channel(grpc_server, options=options) as channel:
        yield document_pb2_grpc.DocumentServiceStub(channel)

def test_process_pdf_integration(grpc_stub):
    # Load a real PDF from the resources directory
    repo_root = Path(__file__).resolve().parent.parent.parent
    pdf_path = repo_root / "resources" / "OriginalApplication.pdf"

    with open(pdf_path, "rb") as f:
        pdf_bytes = f.read()

    request = document_pb2.ProcessPdfRequest(pdf_bytes=pdf_bytes)
    response = grpc_stub.ProcessPdf(request)

    # Verify we get a valid PDF ID (SHA-256 hash usually)
    assert len(response.pdf_id) > 0
    assert response.pdf_id != "mock-pdf-id"

    # Verify we got some extracted values
    # OriginalApplication.pdf should have form fields
    assert len(response.values) > 0

    # Verify registry_json is valid JSON and contains structural data
    import json
    registry = json.loads(response.registry_json)
    assert response.pdf_id in registry
    assert "pages" in registry[response.pdf_id]
    assert "fields" in registry[response.pdf_id]

def test_generate_pdf_integration(grpc_stub):
    # Load a real PDF from the resources directory
    repo_root = Path(__file__).resolve().parent.parent.parent
    pdf_path = repo_root / "resources" / "OriginalApplication.pdf"

    with open(pdf_path, "rb") as f:
        pdf_bytes = f.read()

    # Form data mapping to fields in OriginalApplication.pdf
    form_data = {
        "Full Name": "John Doe",
        "Email": "john.doe@example.com"
    }

    request = document_pb2.GeneratePdfRequest(pdf_bytes=pdf_bytes, form_data=form_data)
    response = grpc_stub.GeneratePdf(request)

    assert len(response.pdf_bytes) > 0
    assert response.pdf_bytes.startswith(b"%PDF-")

def test_generate_docx_integration(grpc_stub):
    # Create a minimal DOCX in memory
    from docx import Document
    from io import BytesIO

    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_table(rows=2, cols=2) # Add a table as the generator expects tables

    docx_stream = BytesIO()
    doc.save(docx_stream)
    docx_bytes = docx_stream.getvalue()

    request = document_pb2.GenerateDocxRequest(docx_bytes=docx_bytes, form_data={"name": "John Doe"})
    response = grpc_stub.GenerateDocx(request)

    assert len(response.docx_bytes) > 0
    # DOCX is a zip file, should start with PK
    assert response.docx_bytes.startswith(b"PK")

def test_stamp_signature_integration(grpc_stub):
    # Load a real PDF from the resources directory
    repo_root = Path(__file__).resolve().parent.parent.parent
    pdf_path = repo_root / "resources" / "OriginalApplication.pdf"

    with open(pdf_path, "rb") as f:
        pdf_bytes = f.read()

    # Create a sample PNG signature image using Pillow
    from PIL import Image
    from io import BytesIO
    img = Image.new('RGBA', (100, 50), color=(255, 0, 0, 128))
    sig_stream = BytesIO()
    img.save(sig_stream, format="PNG")
    sig_bytes = sig_stream.getvalue()

    request = document_pb2.StampSignatureRequest(
        pdf_bytes=pdf_bytes,
        signature_image_bytes=sig_bytes,
        pdf_id="test-pdf-id",
        registry_json="{}",
        cache_mapping_json="{}"
    )
    response = grpc_stub.StampSignature(request)

    assert len(response.pdf_bytes) > 0
    assert response.pdf_bytes.startswith(b"%PDF-")
    # Stamped PDF should be different from original
    assert response.pdf_bytes != pdf_bytes

def test_payload_size_limit(grpc_stub):
    # Create a payload slightly larger than 5MB
    large_payload = b"0" * (5 * 1024 * 1024 + 1)

    # Test ProcessPdf
    request = document_pb2.ProcessPdfRequest(pdf_bytes=large_payload)
    with pytest.raises(grpc.RpcError) as excinfo:
        grpc_stub.ProcessPdf(request)
    assert excinfo.value.code() == grpc.StatusCode.INVALID_ARGUMENT
    assert "Payload exceeds 5MB limit" in excinfo.value.details()

    # Test GeneratePdf
    request = document_pb2.GeneratePdfRequest(pdf_bytes=large_payload, form_data={})
    with pytest.raises(grpc.RpcError) as excinfo:
        grpc_stub.GeneratePdf(request)
    assert excinfo.value.code() == grpc.StatusCode.INVALID_ARGUMENT
    assert "Payload exceeds 5MB limit" in excinfo.value.details()

    # Test GenerateDocx
    request = document_pb2.GenerateDocxRequest(docx_bytes=large_payload, form_data={})
    with pytest.raises(grpc.RpcError) as excinfo:
        grpc_stub.GenerateDocx(request)
    assert excinfo.value.code() == grpc.StatusCode.INVALID_ARGUMENT
    assert "Payload exceeds 5MB limit" in excinfo.value.details()

    # Test StampSignature
    request = document_pb2.StampSignatureRequest(
        pdf_bytes=large_payload,
        signature_image_bytes=b"small",
        pdf_id="test",
        registry_json="{}",
        cache_mapping_json="{}"
    )
    with pytest.raises(grpc.RpcError) as excinfo:
        grpc_stub.StampSignature(request)
    assert excinfo.value.code() == grpc.StatusCode.INVALID_ARGUMENT
    assert "Payload exceeds 5MB limit" in excinfo.value.details()

def test_resource_limits_initialization():
    # Mock resource.setrlimit to verify it is called with the expected limits
    # without actually affecting the test process.
    soft_limit = 1536 * 1024 * 1024
    hard_limit = 2048 * 1024 * 1024

    with patch("resource.setrlimit") as mock_setrlimit:
        configure_resource_limits()
        mock_setrlimit.assert_called_once_with(resource.RLIMIT_AS, (soft_limit, hard_limit))

def test_timeout_error(grpc_stub):
    # Mock process_pdf to sleep for longer than the timeout
    # We patch TIMEOUT_SECONDS to 1 second for a faster test
    with patch("src.server.TIMEOUT_SECONDS", 1):
        with patch("src.server.process_pdf") as mock_process:
            def slow_process(*args, **kwargs):
                time.sleep(2)
                return "id", {}, {}
            mock_process.side_effect = slow_process

            request = document_pb2.ProcessPdfRequest(pdf_bytes=b"fake-pdf")
            with pytest.raises(grpc.RpcError) as excinfo:
                grpc_stub.ProcessPdf(request)

            assert excinfo.value.code() == grpc.StatusCode.DEADLINE_EXCEEDED
            assert "Processing timed out" in excinfo.value.details()
