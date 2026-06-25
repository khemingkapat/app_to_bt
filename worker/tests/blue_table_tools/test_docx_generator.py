import os
import docx
import pytest
from datetime import datetime
from src.blue_table_tools.docx_generator import calculate_age, fill_blue_table_docx, apply_acceptance_rules, resolve_plan_combination

def test_calculate_age():
    # Test DD/MM/YYYY
    current_year = datetime.today().year
    assert calculate_age(f"15/08/1990") == str(current_year - 1990 - ((datetime.today().month, datetime.today().day) < (8, 15)))
    
    # Test YYYY-MM-DD
    assert calculate_age("1985-12-01") == str(current_year - 1985 - ((datetime.today().month, datetime.today().day) < (12, 1)))
    
    # Test regex fallback (only year)
    assert calculate_age("Born in 2005") == str(current_year - 2005)
    
    # Empty case
    assert calculate_age("") == ""
    assert calculate_age(None) == ""

def test_fill_blue_table_docx():
    template_path = "./resources/BlueTable.docx"
    if not os.path.exists(template_path):
        pytest.skip("BlueTable.docx template not found")
        
    sample_data = {
        "name": "Alice Wonderland",
        "dob": "10/10/1992",
        "id_card_no": "9999999999999",
        "nationality": "Wonderlander",
        "beneficiary": "Bob Wonderland",
        "bene_relation": "Sibling",
        "occupation": "Dreamer",
        "agent": "AG-1234",
        "plan": "Plan 4 (10M)",
        "deductible": "0 THB",
        "premium": "35,000 THB",
        "effective_date": "15/06/2026",
        "present_address": "Rabbit Hole 1",
        "tel": "0800000000",
        "email": "alice@wonderland.com",
        "sp_name": "Bob Wonderland",
        "sp_dob": "01/01/1990",
        "c1_name": "Charlie Wonderland",
        "c1_dob": "05/05/2018",
    }
    
    # Generate filled docx stream
    stream = fill_blue_table_docx(template_path, sample_data)
    assert stream is not None
    assert len(stream.getvalue()) > 0
    
    # Parse back the docx from stream and check values
    stream.seek(0)
    doc = docx.Document(stream)
    
    # Check Table 0 (Main Insured) values
    t0 = doc.tables[0]
    t0_vals = {row.cells[0].text.strip(): row.cells[1].text.strip() for row in t0.rows}
    
    assert t0_vals.get("Main Insured") == "Alice Wonderland"
    assert t0_vals.get("Date of Birth") == "10/10/1992"
    assert t0_vals.get("ID No./Passport No.") == "9999999999999"
    assert t0_vals.get("Nationality") == "Wonderlander"
    assert t0_vals.get("Beneficiary name") == "Bob Wonderland"
    assert t0_vals.get("Relation") == "Sibling"
    assert t0_vals.get("Occupation") == "Dreamer"
    assert t0_vals.get("Agent CODE/Name") == "AG-1234"
    assert t0_vals.get("Plan") == "Plan 4 (10M)"
    assert t0_vals.get("Deductible") == "0 THB"
    assert t0_vals.get("Premium") == "35,000 THB"
    assert t0_vals.get("Effective date") == "15/06/2026"
    assert t0_vals.get("Personal Address") == "Rabbit Hole 1"
    assert t0_vals.get("Tel") == "0800000000"
    assert t0_vals.get("Email") == "alice@wonderland.com"
    
    # Check age calculation was filled
    expected_age = calculate_age("10/10/1992")
    assert t0_vals.get("Age") == expected_age
    
    # Check Table 1 (Spouse) values
    t1 = doc.tables[1]
    t1_vals = {row.cells[0].text.strip(): row.cells[1].text.strip() for row in t1.rows}
    assert t1_vals.get("Spouse") == "Bob Wonderland"
    assert t1_vals.get("Date of Birth") == "01/01/1990"
    assert t1_vals.get("Age") == calculate_age("01/01/1990")
    
    # Check Table 2 (Child 1) values
    t2 = doc.tables[2]
    t2_vals = {row.cells[0].text.strip(): row.cells[1].text.strip() for row in t2.rows}
    assert t2_vals.get("Child # 1") == "Charlie Wonderland"
    assert t2_vals.get("Date of Birth") == "05/05/2018"
    assert t2_vals.get("Age") == calculate_age("05/05/2018")

def test_apply_acceptance_rules():
    # Clean case
    data = {
        "name": "Alice",
        "exclusions": "None",
        "sp_name": "Bob",
        "sp_exclusions": "",
        "c1_name": "",
        "c1_exclusions": "Mild Asthma",
    }
    res = apply_acceptance_rules(data)
    assert res["acceptance_conditions"] == "Accepted"
    assert res["sp_acceptance_conditions"] == "Accepted"
    assert res["c1_acceptance_conditions"] == ""

    # Unclean case
    data = {
        "name": "Alice",
        "exclusions": "Mild Asthma",
        "sp_name": "Bob",
        "sp_exclusions": "None.",
    }
    res = apply_acceptance_rules(data)
    assert res["acceptance_conditions"] == "Accepted with exclusion"
    assert res["sp_acceptance_conditions"] == "Accepted"

def test_resolve_plan_combination():
    from src.pdf_processor.inverter import load_config_by_pdf_id
    config = load_config_by_pdf_id(None)
    combo_map = config.get("combinations_map", {})
    has_mocka = any(k.startswith("MOCKA") for k in combo_map.keys())
    prefix_a = "MOCKA" if has_mocka else "ESSENTIAL"
    prefix_b = "MOCKB" if has_mocka else "VISA"
    
    # Test valid combination lookup (IPD only)
    data = {
        "plan": "Plan 2-IPD",
        "deductible": "20k"
    }
    resolved = resolve_plan_combination(data)
    assert resolved["plan"] == f"{prefix_a}2-IPD DD 20,000 (127)"
    assert resolved["deductible"] == "20,000"

    # Test reverse order combination lookup
    data = {
        "plan": "IPD-Plan 2",
        "deductible": "20k"
    }
    resolved = resolve_plan_combination(data)
    assert resolved["plan"] == f"{prefix_a}2-IPD DD 20,000 (127)"
    assert resolved["deductible"] == "20,000"

    # Test IPD+OPD combination lookup with specific OPD limits
    data = {
        "plan": "Plan 1-IPD+OPD-3k * 30 times / year",
        "deductible": "20k"
    }
    resolved = resolve_plan_combination(data)
    assert resolved["plan"] == f"{prefix_a}1-IPD+OPD(3k * 30 times / year) DD 20,000 (107)"
    assert resolved["deductible"] == "20,000"

    # Test already resolved plan (should skip or re-resolve to same)
    data = {
        "plan": f"{prefix_a}2-IPD DD 20,000 (127)",
        "deductible": "20,000"
    }
    resolved = resolve_plan_combination(data)
    assert resolved["plan"] == f"{prefix_a}2-IPD DD 20,000 (127)"

    # Test already resolved plan with a changed deductible (should re-resolve dynamically)
    data = {
        "plan": f"{prefix_a}2-IPD DD 20,000 (127)",
        "deductible": "0"
    }
    resolved = resolve_plan_combination(data)
    assert resolved["plan"] == f"{prefix_a}2-IPD DD 0 (126)"
    assert resolved["deductible"] == "0"

    # Test EasyCare Visa combination lookup (legacy format)
    data = {
        "plan": "VISA Plan 1",
        "deductible": "100k"
    }
    resolved = resolve_plan_combination(data)
    assert resolved["plan"] == f"{prefix_b}1 DD 100,000 (201)"
    assert resolved["deductible"] == "100,000"

    # Test new direct product plan prefix format (ESSENTIAL and VISA)
    data = {
        "plan": f"{prefix_a}2-IPD",
        "deductible": "20k"
    }
    resolved = resolve_plan_combination(data)
    assert resolved["plan"] == f"{prefix_a}2-IPD DD 20,000 (127)"

    data = {
        "plan": f"{prefix_b}1",
        "deductible": "100k"
    }
    resolved = resolve_plan_combination(data)
    assert resolved["plan"] == f"{prefix_b}1 DD 100,000 (201)"


def test_policy_version_and_general_conditions():
    template_path = "./resources/BlueTable.docx"
    if not os.path.exists(template_path):
        pytest.skip("BlueTable.docx template not found")
        
    data_ess = {
        "name": "John Doe",
        "product_name": "ESSENTIAL",
        "policy_version": "Thai",
    }
    
    stream_ess = fill_blue_table_docx(template_path, data_ess)
    doc_ess = docx.Document(stream_ess)
    
    t0 = doc_ess.tables[0]
    t0_vals = {row.cells[0].text.strip(): row.cells[1].text.strip() for row in t0.rows}
    t0_labels = [row.cells[0].text.strip() for row in t0.rows]
    
    assert t0_vals.get("Policy Version") == "TH"
    gen_row_label = next(l for l in t0_labels if "General Conditions" in l)
    assert gen_row_label == "General Conditions:\nMockCare Plan A"
    
    # Check that cell 1 has correct general conditions text from config
    from src.pdf_processor.inverter import load_config_by_pdf_id
    config = load_config_by_pdf_id(None)
    expected_msg_ess = config.get("general_conditions", {}).get("MockCare Plan A", {}).get("Thai", "")
    assert t0_vals.get(gen_row_label) == expected_msg_ess

    data_visa = {
        "name": "John Doe",
        "product_name": "EASYCARE",
        "policy_version": "English",
    }
    
    stream_visa = fill_blue_table_docx(template_path, data_visa)
    doc_visa = docx.Document(stream_visa)
    
    t0_visa = doc_visa.tables[0]
    t0_visa_vals = {row.cells[0].text.strip(): row.cells[1].text.strip() for row in t0_visa.rows}
    t0_visa_labels = [row.cells[0].text.strip() for row in t0_visa.rows]
    
    assert t0_visa_vals.get("Policy Version") == "EN"
    gen_row_label_visa = next(l for l in t0_visa_labels if "General Conditions" in l)
    assert gen_row_label_visa == "General Conditions:\nMockCare Plan B"
    
    expected_msg_visa = config.get("general_conditions", {}).get("MockCare Plan B", {}).get("English", "")
    assert t0_visa_vals.get(gen_row_label_visa) == expected_msg_visa
