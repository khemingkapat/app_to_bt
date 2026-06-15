import re
from io import BytesIO
from datetime import datetime
import docx

def calculate_age(dob_str: str) -> str:
    """
    Calculates age based on Date of Birth string.
    Supports formats: DD/MM/YYYY, YYYY-MM-DD, DD-MM-YYYY, YYYY/MM/DD.
    Falls back to simple birth-year subtraction if format is unparseable.
    """
    if not dob_str:
        return ""
    
    # Try common datetime formats
    for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%Y/%m/%d"):
        try:
            dob_dt = datetime.strptime(dob_str.strip(), fmt)
            today = datetime.today()
            age = today.year - dob_dt.year - ((today.month, today.day) < (dob_dt.month, dob_dt.day))
            return str(age)
        except ValueError:
            continue
            
    # Regex fallback to find a 4-digit year (e.g. 1990)
    years = re.findall(r"\b(19\d\d|20\d\d)\b", dob_str)
    if years:
        dob_year = int(years[0])
        return str(datetime.today().year - dob_year)
        
    return ""

def apply_acceptance_rules(data: dict) -> dict:
    """
    Automatically populates the 'Acceptance Conditions' for each member block
    based on the presence of exclusions.
    - Clean case: 'Accepted'
    - Case with exclusions: 'Accepted with exclusion'
    """
    updated = data.copy()
    blocks = [
        ("name", "exclusions", "acceptance_conditions"),
        ("sp_name", "sp_exclusions", "sp_acceptance_conditions"),
        ("c1_name", "c1_exclusions", "c1_acceptance_conditions"),
        ("c2_name", "c2_exclusions", "c2_acceptance_conditions"),
        ("c3_name", "c3_exclusions", "c3_acceptance_conditions"),
    ]
    
    # Values of exclusions that indicate a "clean" case (no exclusions)
    clean_indicators = {
        "", "none", "none.", "n/a", "na", "no", "no exclusion", 
        "no exclusions", "-", "clean", "nil", "no pre-existing conditions"
    }
    
    for name_key, excl_key, status_key in blocks:
        name_val = str(updated.get(name_key) or "").strip()
        if name_val:
            excl_val = str(updated.get(excl_key) or "").strip()
            # Remove trailing/leading spaces and lowercase for check
            check_val = excl_val.lower().rstrip(".")
            if not check_val or check_val in clean_indicators:
                updated[status_key] = "Accepted"
            else:
                updated[status_key] = "Accepted with exclusion"
        else:
            # If the person doesn't exist, clear acceptance conditions
            updated[status_key] = ""
            
    return updated

def resolve_plan_combination(data: dict) -> dict:
    """
    Resolves the combined plan name and code from raw mapping fields
    and updates the 'plan' and 'deductible' fields in the data dict.
    """
    updated = data.copy()
    plan_val = str(updated.get("plan") or "").strip()
    ded_val = str(updated.get("deductible") or "").strip()
    
    if "(" in plan_val and ")" in plan_val:
        return updated
        
    parts = [p.strip() for p in plan_val.split("-") if p.strip()]
    plan_tier = ""
    optional_benefit = ""
    opd_choice = ""
    
    valid_benefits = ("IPD", "IPD+OPD", "IPD+OPD+WELLNESS")
    valid_opd_choices = ("3k * 30 times / year", "50k per year")
    
    for p in parts:
        if "Plan" in p:
            plan_tier = p.replace("Plan", "").strip()
        elif p in valid_benefits:
            optional_benefit = p
        elif p in valid_opd_choices:
            opd_choice = p
            
    deductible_amount = ded_val.replace("k", ",000")
    if deductible_amount == "0,000":
        deductible_amount = "0"
         
    if plan_tier and optional_benefit:
        # Build the combo key depending on benefit type
        if optional_benefit == "IPD":
            combo_key = f"ESSENTIAL{plan_tier}-IPD DD {deductible_amount}"
        else:
            combo_key = f"ESSENTIAL{plan_tier}-{optional_benefit}({opd_choice}) DD {deductible_amount}"
            
        try:
            from src.pdf_processor.inverter import load_product_config
            config = load_product_config("./config/health_and_accident.json")
            combo_map = config.get("combinations_map", {})
            plan_code = combo_map.get(combo_key)
            if plan_code:
                updated["plan"] = f"{combo_key} ({plan_code})"
                updated["deductible"] = deductible_amount
        except Exception:
            pass
            
    return updated


def fill_blue_table_docx(template_path: str, data: dict) -> BytesIO:
    """
    Fills the BlueTable.docx template tables with the provided data dict.
    Returns the filled file as a BytesIO stream.
    """
    data = resolve_plan_combination(data)
    data = apply_acceptance_rules(data)
    doc = docx.Document(template_path)
    
    # Table 0: Main Insured
    age = data.get("age", "")
    if not age and data.get("dob"):
        age = calculate_age(data["dob"])
        
    t0_mapping = {
        "Main Insured": data.get("name", ""),
        "Date of Birth": data.get("dob", ""),
        "Age": age,
        "ID No./Passport No.": data.get("id_card_no", ""),
        "Nationality": data.get("nationality", ""),
        "Beneficiary name": data.get("beneficiary", ""),
        "Relation": data.get("bene_relation", ""),
        "Occupation": data.get("occupation", ""),
        "Agent CODE/Name": data.get("agent", ""),
        "Plan": data.get("plan", ""),
        "Deductible": data.get("deductible", ""),
        "Premium": data.get("premium", ""),
        "Effective date": data.get("effective_date", ""),
        "Personal Address": data.get("present_address", ""),
        "Tel": data.get("tel", ""),
        "Email": data.get("email", ""),
        "Payor Name": data.get("payor_name", ""),
        "Payor Address": data.get("payor_address", ""),
        "TAX ID": data.get("tax_id", ""),
        "Acceptance Conditions": data.get("acceptance_conditions", ""),
        "Exclusions": data.get("exclusions", ""),
    }

    # Table 1: Spouse
    sp_age = data.get("sp_age", "")
    if not sp_age and data.get("sp_dob"):
        sp_age = calculate_age(data["sp_dob"])
        
    t1_mapping = {
        "Spouse": data.get("sp_name", ""),
        "Date of Birth": data.get("sp_dob", ""),
        "Age": sp_age,
        "ID No./Passport No.": data.get("sp_id_card_no", ""),
        "Nationality": data.get("sp_nationality", ""),
        "Beneficiary name": data.get("sp_beneficiary", ""),
        "Relation": data.get("sp_bene_relation", ""),
        "Occupation": data.get("sp_occupation", ""),
        "Acceptance Conditions": data.get("sp_acceptance_conditions", ""),
        "Exclusions": data.get("sp_exclusions", ""),
    }

    # Table 2: Child 1
    c1_age = data.get("c1_age", "")
    if not c1_age and data.get("c1_dob"):
        c1_age = calculate_age(data["c1_dob"])
        
    t2_mapping = {
        "Child # 1": data.get("c1_name", ""),
        "Date of Birth": data.get("c1_dob", ""),
        "Age": c1_age,
        "ID No./Passport No.": data.get("c1_id_card_no", ""),
        "Nationality": data.get("c1_nationality", ""),
        "Beneficiary name": data.get("c1_beneficiary", ""),
        "Relation": data.get("c1_bene_relation", ""),
        "Occupation": data.get("c1_occupation", ""),
        "Acceptance Conditions": data.get("c1_acceptance_conditions", ""),
        "Exclusions": data.get("c1_exclusions", ""),
    }

    # Table 3: Child 2
    c2_age = data.get("c2_age", "")
    if not c2_age and data.get("c2_dob"):
        c2_age = calculate_age(data["c2_dob"])
        
    t3_mapping = {
        "Child # 2": data.get("c2_name", ""),
        "Date of Birth": data.get("c2_dob", ""),
        "Age": c2_age,
        "ID No./Passport No.": data.get("c2_id_card_no", ""),
        "Nationality": data.get("c2_nationality", ""),
        "Beneficiary name": data.get("c2_beneficiary", ""),
        "Relation": data.get("c2_bene_relation", ""),
        "Occupation": data.get("c2_occupation", ""),
        "Acceptance Conditions": data.get("c2_acceptance_conditions", ""),
        "Exclusions": data.get("c2_exclusions", ""),
    }

    # Table 4: Child 3
    c3_age = data.get("c3_age", "")
    if not c3_age and data.get("c3_dob"):
        c3_age = calculate_age(data["c3_dob"])
        
    t4_mapping = {
        "Child # 3": data.get("c3_name", ""),
        "Date of Birth": data.get("c3_dob", ""),
        "Age": c3_age,
        "ID No./Passport No.": data.get("c3_id_card_no", ""),
        "Nationality": data.get("c3_nationality", ""),
        "Relation": data.get("c3_bene_relation", ""),
        "Occupation": data.get("c3_occupation", ""),
        "Acceptance Conditions": data.get("c3_acceptance_conditions", ""),
        "Exclusions": data.get("c3_exclusions", ""),
    }

    mappings = [t0_mapping, t1_mapping, t2_mapping, t3_mapping, t4_mapping]

    def set_cell_text(cell, text):
        p = cell.paragraphs[0]
        p.text = str(text) if text is not None else ""
        # Remove extra paragraphs in the cell to avoid adding unintended newlines
        for extra_p in cell.paragraphs[1:]:
            p_element = extra_p._p
            p_element.getparent().remove(p_element)

    # Populate matching row keys
    for i, table in enumerate(doc.tables):
        if i >= len(mappings):
            break
        mapping = mappings[i]
        
        for row in table.rows:
            col0_text = row.cells[0].text.strip()
            if col0_text in mapping:
                set_cell_text(row.cells[1], mapping[col0_text])

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
