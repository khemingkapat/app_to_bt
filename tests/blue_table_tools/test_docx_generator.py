import os
import docx
import pytest
from datetime import datetime
from src.blue_table_tools.docx_generator import calculate_age, fill_blue_table_docx

def test_calculate_age():
    # Test DD/MM/YYYY
    current_year = datetime.today().year
    assert calculate_age(f"15/08/1990") == str(current_year - 1990 - ((datetime.today().month, datetime.today().day) < (8, 15)))
    
    # Test YYYY-MM-DD
    assert calculate_age("1985-12-01") == str(current_year - 1985 - ((datetime.today().month, datetime.today().day) < (12, 1)))
    
    # Test regex fallback (only year)
    assert calculate_age("Born in 2005") == str(current_year - 2005)
    
    # Empty case
    assert calculate_age("") == ""
    assert calculate_age(None) == ""

def test_fill_blue_table_docx():
    template_path = "./resources/BlueTable.docx"
    if not os.path.exists(template_path):
        pytest.skip("BlueTable.docx template not found")
        
    sample_data = {
        "name": "Alice Wonderland",
        "dob": "10/10/1992",
        "id_card_no": "9999999999999",
        "nationality": "Wonderlander",
        "beneficiary": "Bob Wonderland",
        "bene_relation": "Sibling",
        "occupation": "Dreamer",
        "agent": "AG-1234",
        "plan": "Plan 4 (10M)",
        "deductible": "0 THB",
        "premium": "35,000 THB",
        "effective_date": "15/06/2026",
        "present_address": "Rabbit Hole 1",
        "tel": "0800000000",
        "email": "alice@wonderland.com",
        "sp_name": "Bob Wonderland",
        "sp_dob": "01/01/1990",
        "c1_name": "Charlie Wonderland",
        "c1_dob": "05/05/2018",
    }
    
    # Generate filled docx stream
    stream = fill_blue_table_docx(template_path, sample_data)
    assert stream is not None
    assert len(stream.getvalue()) > 0
    
    # Parse back the docx from stream and check values
    stream.seek(0)
    doc = docx.Document(stream)
    
    # Check Table 0 (Main Insured) values
    t0 = doc.tables[0]
    t0_vals = {row.cells[0].text.strip(): row.cells[1].text.strip() for row in t0.rows}
    
    assert t0_vals.get("Main Insured") == "Alice Wonderland"
    assert t0_vals.get("Date of Birth") == "10/10/1992"
    assert t0_vals.get("ID No./Passport No.") == "9999999999999"
    assert t0_vals.get("Nationality") == "Wonderlander"
    assert t0_vals.get("Beneficiary name") == "Bob Wonderland"
    assert t0_vals.get("Relation") == "Sibling"
    assert t0_vals.get("Occupation") == "Dreamer"
    assert t0_vals.get("Agent CODE/Name") == "AG-1234"
    assert t0_vals.get("Plan") == "Plan 4 (10M)"
    assert t0_vals.get("Deductible") == "0 THB"
    assert t0_vals.get("Premium") == "35,000 THB"
    assert t0_vals.get("Effective date") == "15/06/2026"
    assert t0_vals.get("Personal Address") == "Rabbit Hole 1"
    assert t0_vals.get("Tel") == "0800000000"
    assert t0_vals.get("Email") == "alice@wonderland.com"
    
    # Check age calculation was filled
    expected_age = calculate_age("10/10/1992")
    assert t0_vals.get("Age") == expected_age
    
    # Check Table 1 (Spouse) values
    t1 = doc.tables[1]
    t1_vals = {row.cells[0].text.strip(): row.cells[1].text.strip() for row in t1.rows}
    assert t1_vals.get("Spouse") == "Bob Wonderland"
    assert t1_vals.get("Date of Birth") == "01/01/1990"
    assert t1_vals.get("Age") == calculate_age("01/01/1990")
    
    # Check Table 2 (Child 1) values
    t2 = doc.tables[2]
    t2_vals = {row.cells[0].text.strip(): row.cells[1].text.strip() for row in t2.rows}
    assert t2_vals.get("Child # 1") == "Charlie Wonderland"
    assert t2_vals.get("Date of Birth") == "05/05/2018"
    assert t2_vals.get("Age") == calculate_age("05/05/2018")
